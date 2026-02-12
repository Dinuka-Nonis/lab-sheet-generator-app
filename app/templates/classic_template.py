"""
Classic Template for Lab Sheet Generator
Original template design with colored bar and centered layout
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import os

from app.core.template_manager import BaseTemplate, register_template


class ClassicTemplate(BaseTemplate):
    """Classic lab sheet template with blue colored bar."""
    
    template_id = "classic"
    template_name = "Classic Template"
    template_description = "Original design with blue header bar and centered layout"
    
    def get_required_fonts(self):
        """Classic template uses Times New Roman (system font)."""
        return ["Times New Roman"]
    
    def add_colored_rectangle(self, doc):
        """Add a blue colored rectangle at the top of the document."""
        bar = doc.add_paragraph()
        bar_format = bar.paragraph_format
        bar_format.space_before = Pt(0)
        bar_format.space_after = Pt(0)
        bar_format.line_spacing = Pt(15)
        bar_format.left_indent = Inches(-0.5)
        bar_format.right_indent = Inches(-0.5)
        
        # Set shading (background color) for the paragraph
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '156082')  # Blue color
        bar._element.get_or_add_pPr().append(shading_elm)
        
        return bar
    
    def generate(self, student_name, student_id, module_name, module_code, 
                 sheet_label, logo_path=None):
        """
        Generate a lab sheet using the classic template.
        
        Args:
            student_name: Student's full name
            student_id: Student ID number
            module_name: Name of the module
            module_code: Module code
            sheet_label: Sheet label (e.g., "Practical 06")
            logo_path: Path to the university logo image
            
        Returns:
            str: Filename of generated document
        """
        # Create a new Document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Add the blue colored rectangle at the top
        self.add_colored_rectangle(doc)
        
        # Add the university logo (centered from page edges)
        if logo_path and Path(logo_path).exists():
            logo_paragraph = doc.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_paragraph.paragraph_format.left_indent = Inches(-0.5)
            logo_paragraph.paragraph_format.right_indent = Inches(-0.5)
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(str(logo_path), width=Inches(1.1), height=Inches(1.05))
            logo_paragraph.paragraph_format.space_after = Pt(6)
        
        # Add module name and code (centered, size 20)
        module_paragraph = doc.add_paragraph()
        module_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        module_paragraph.paragraph_format.left_indent = Inches(-0.5)
        module_paragraph.paragraph_format.right_indent = Inches(-0.5)
        module_run = module_paragraph.add_run(f'{module_name} -- {module_code}')
        module_run.bold = True
        module_run.font.size = Pt(20)
        module_run.font.name = 'Times New Roman'
        module_paragraph.paragraph_format.space_after = Pt(12)
        
        # Add sheet label (size 12, left aligned)
        sheet_paragraph = doc.add_paragraph()
        sheet_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sheet_run = sheet_paragraph.add_run(f'{sheet_label}')
        sheet_run.bold = True
        sheet_run.font.size = Pt(12)
        sheet_run.font.name = 'Times New Roman'
        sheet_paragraph.paragraph_format.space_after = Pt(0)
        
        # Add student name and ID on next line (size 12, left aligned)
        name_paragraph = doc.add_paragraph()
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_paragraph.add_run(f'{student_name} - {student_id}')
        name_run.bold = True
        name_run.font.size = Pt(12)
        name_run.font.name = 'Times New Roman'
        name_paragraph.paragraph_format.space_after = Pt(6)
        
        # Add the horizontal line
        line_paragraph = doc.add_paragraph()
        line_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        line_run = line_paragraph.add_run('_' * 95)
        line_run.font.name = 'Times New Roman'
        line_paragraph.paragraph_format.space_after = Pt(12)
        
        # Save the document
        output_filename = f'{sheet_label.replace(" ", "_")}_{student_id}.docx'
        doc.save(output_filename)
        
        return output_filename


# Register the template
register_template(ClassicTemplate)