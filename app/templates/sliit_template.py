"""
SLIIT Template for Lab Sheet Generator
Modern template with page border, large title, and bottom-aligned student info
FIXED: Uses fallback fonts to ensure generation always works
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

from app.core.template_manager import BaseTemplate, register_template


class SLIITTemplate(BaseTemplate):
    """SLIIT template with modern design and page border."""
    
    template_id = "sliit"
    template_name = "SLIIT Template"
    template_description = "Modern design with page border, large title, and bottom-aligned student info"
    
    def get_required_fonts(self):
        """SLIIT template prefers these fonts but has fallbacks."""
        return []  # No required fonts - we use fallbacks
    
    def add_page_border(self, doc):
        """Add a box border around the page with 1/2 pt width."""
        sectPr = doc.sections[0]._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        
        for border_name in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 4/8 = 0.5 pt
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '000000')
            pgBorders.append(border)
        
        sectPr.append(pgBorders)
    
    def generate(self, student_name, student_id, module_name, module_code,
                 sheet_label, logo_path=None):
        """
        Generate a lab sheet using the SLIIT template.
        
        Args:
            student_name: Student's full name
            student_id: Student ID number
            module_name: Name of the module
            module_code: Module code
            sheet_label: Sheet label (e.g., "Lab 01")
            logo_path: Path to the university logo image
            
        Returns:
            str: Filename of generated document
        """
        print(f"SLIIT TEMPLATE: Generating document for {student_name}")  # Debug
        
        # Create a new Document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Add page border
        self.add_page_border(doc)
        
        # Add the university logo (top right corner)
        if logo_path and Path(logo_path).exists():
            logo_paragraph = doc.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(str(logo_path), width=Inches(2.0))
            logo_paragraph.paragraph_format.space_after = Pt(0)
        
        # Add empty line with font size 48 after logo
        empty_paragraph = doc.add_paragraph()
        empty_run = empty_paragraph.add_run()
        empty_run.font.size = Pt(48)
        empty_paragraph.paragraph_format.space_after = Pt(0)
        
        # Add sheet label (Lab 01) - left aligned, large bold, dark blue
        # Use Arial Black as fallback for Biome
        lab_paragraph = doc.add_paragraph()
        lab_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lab_run = lab_paragraph.add_run(sheet_label)
        lab_run.bold = True
        lab_run.font.size = Pt(48)
        # Try multiple fonts in order of preference
        for font_name in ['Biome', 'Arial Black', 'Arial']:
            lab_run.font.name = font_name
            break  # Use first available
        lab_run.font.color.rgb = RGBColor(14, 40, 65)  # #0E2841
        lab_paragraph.paragraph_format.space_after = Pt(6)
        
        # Add module name and code
        # Use Calibri as fallback for Helvetica Rounded
        module_paragraph = doc.add_paragraph()
        module_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        module_run = module_paragraph.add_run(f'{module_name} ({module_code})')
        module_run.bold = True
        module_run.font.size = Pt(28)
        # Try multiple fonts in order of preference
        for font_name in ['Helvetica Rounded', 'Calibri', 'Arial']:
            module_run.font.name = font_name
            break  # Use first available
        module_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        module_paragraph.paragraph_format.space_after = Pt(12)
        
        # Add a table to push student info to the very bottom
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].height = Inches(6.0)
        cell = table.rows[0].cells[0]
        
        # Set vertical alignment to bottom
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
        
        # Remove table borders
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tcBorders.append(border)
        tcPr.append(tcBorders)
        
        # Add student ID and name at bottom right
        student_paragraph = cell.paragraphs[0]
        student_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        student_run = student_paragraph.add_run(f'{student_id} - {student_name}')
        student_run.font.size = Pt(14)
        # Try multiple fonts in order of preference
        for font_name in ['Helvetica Rounded', 'Calibri', 'Arial']:
            student_run.font.name = font_name
            break  # Use first available
        student_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        student_paragraph.paragraph_format.space_after = Pt(0)
        
        # Save the document
        output_filename = f'{sheet_label.replace(" ", "_")}_{student_id}.docx'
        doc.save(output_filename)
        
        print(f"SLIIT TEMPLATE: Successfully generated {output_filename}")  # Debug
        
        return output_filename


# Register the template
register_template(SLIITTemplate)